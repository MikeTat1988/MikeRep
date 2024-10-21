import os
from docx import Document
from docx.shared import Pt

def add_files_from_folder(doc, folder_path, folder_name, file_counter):
    """Add files from a folder with their content and update the file counter."""
    print(f"Scanning folder: {folder_path}")  # Debugging
    for root, _, files in os.walk(folder_path):
        relative_folder = os.path.relpath(root, folder_path)
        folder_display = folder_name if relative_folder == '.' else f"{folder_name}\\{relative_folder}"

        for file in files:
            file_path = os.path.join(root, file)
            print(f"Found file: {file_path}")  # Debugging
            add_file_heading(doc, folder_display, file)
            add_file_content(doc, file_path)
            file_counter[0] += 1

def add_file_heading(doc, folder_display, file):
    """Add the file name as bold and underlined heading."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{folder_display}\\{file}:")
    run.bold = True
    run.underline = True
    paragraph.style = 'Normal'

def add_file_content(doc, file_path):
    """Add the content of a file to the document without extra spaces."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.read().splitlines()  # Avoid adding extra spaces
            for line in lines:
                doc.add_paragraph(line, style='NoSpacing')
    except Exception as e:
        doc.add_paragraph(f"[Error reading {file_path}: {e}]")

def add_files_list(doc, folder_path, folder_name, file_counter):
    """List files from a specific folder without adding their content."""
    if os.path.exists(folder_path):
        doc.add_heading(f"Files in {folder_name}:", level=2)
        for file in os.listdir(folder_path):
            file_path = os.path.join(folder_path, file)
            if os.path.isfile(file_path):
                doc.add_paragraph(file, style='NoSpacing')
                file_counter[0] += 1

def generate_document():
    """Generate a Word document with the project structure and file contents."""
    doc = Document()

    # Paths
    main_folder = r'C:\Users\micha\MikeRep\ePicSearch\ePicSearch\ePicSearch'
    core_folder = r'C:\Users\micha\MikeRep\ePicSearch\ePicSearch\ePicSearch.Core'
    output_file = r'C:\Users\micha\MikeRep\Project_Structure.docx'

    # File counter to keep track of total files
    file_counter = [0]

    # Write the project structure and files
    doc.add_heading('Project Structure', 0)

    # Main project
    doc.add_heading('Project: ePicSearch', level=1)
    main_files = [
        "ePicSearch.csproj", "App.xaml","App.xaml.cs", "AppShell.xaml",
        "AppShell.xaml.cs", "MauiProgram.cs"
    ]
    for file in main_files:
        file_path = os.path.join(main_folder, file)
        print(f"Checking file: {file_path}")  # Debugging
        if os.path.exists(file_path):
            add_file_heading(doc, '', file)
            add_file_content(doc, file_path)
            file_counter[0] += 1
        else:
            print(f"File not found: {file_path}")  # Debugging

    for folder in ["Entities", "Services", "Views", "Behaviors"]:
        folder_path = os.path.join(main_folder, folder)
        if os.path.exists(folder_path):
            add_files_from_folder(doc, folder_path, folder, file_counter)
        else:
            print(f"Folder not found: {folder_path}")  # Debugging

    # List files from Resources/Images
    images_folder = os.path.join(main_folder, "Resources", "Images")
    add_files_list(doc, images_folder, "Resources\\Images", file_counter)

    # Core project
    doc.add_heading('Project: ePicSearch.Core', level=1)
    core_proj_file = os.path.join(core_folder, "ePicSearch.Infrastructure.csproj")
    if os.path.exists(core_proj_file):
        add_file_heading(doc, '', "ePicSearch.Infrastructure.csproj")
        add_file_content(doc, core_proj_file)
        file_counter[0] += 1

    for folder in ["Entities", "Services"]:
        folder_path = os.path.join(core_folder, folder)
        if os.path.exists(folder_path):
            add_files_from_folder(doc, folder_path, folder, file_counter)

    # Write total file count at the beginning
    doc.add_paragraph(f"Total files: {file_counter[0]}", style='Heading 1')

    # Save the document
    doc.save(output_file)
    print(f"Document saved at {output_file}")

if __name__ == "__main__":
    generate_document()
